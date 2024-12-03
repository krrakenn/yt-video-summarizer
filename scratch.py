from transformers import pipeline
summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=-1)  # Use CPU

def chunk_text(text, max_length=512):
    words = text.split()
    for i in range(0, len(words), max_length):
        yield ' '.join(words[i:i + max_length])

text = """- When QR codes first came
out, I thought they were awful, never going to catch on. This is a flow chart from the time that really resonated with me. The problem as I saw it
was that QR codes are ugly, and they mean nothing to people. I would rather just see a website or a word that I could Google. QR codes are a language for
machines, and I am a human. But I was wrong. QR codes obviously turned
out to be so useful that they are now ubiquitous, used in everything from tickets to restaurant menus and advertising. In some countries, they're the most common
way to exchange money. And the story of QR codes
is a very human one. The origin of these checkerboard
patterns actually dates back to our first efforts
at digitizing information. In 1825, there was a renowned painter who lived in New Haven, Connecticut, with his wife and two kids. His big break came one
day when he was invited to paint a portrait of
the Marquis de Lafayette, a hero of the American Revolution. Even though his wife was expecting their third child any day, the opportunity was too good to pass up, and he hastily set off for Washington DC where Lafayette was waiting. There the painter wrote to his wife describing his
first meeting with Lafayette, signing off with the words, "Will write again soon. Love to all the children. In the greatest haste, but with the same ardent
affection as ever, thy loving husband." After a few days with no reply, a courier delivered a letter which said his wife was
ill after childbirth. Worried, the painter rushed home. He traveled by horse and
wagon, day and night, managing to arrive back in
New Haven in several days. (melancholic music) But it was far too late. His wife had died. Not only that, he had missed the funeral. Her body was already buried in the ground. The painter's name was
Samuel Finley Breese Morse. (inquisitive music) From that day forward, Morse
set out to find a faster way to communicate over long distances. He got a job at New York University where he attended lectures on electricity, a rapidly developing field at the time. In 1836, along with Joseph
Henry and Alfred Vail, he devised a machine that could send electrical
pulses along a wire. This was not the first electric telegraph, but it was the simplest. In the UK, another team had
set up a series of circuits to move five magnetic needles to point at letters and numbers. Morse's system required
only a single circuit, but the simplicity of the apparatus demanded a cleverer method
of encoding information. On the circuit, you could
send short or long pulses. Morse turned these into dots and dashes. The most common letters could be sent with a single key press, a dot for an E and a dash for a T. The other letters were
arranged by frequency and assigned increasingly complex codes. These symbols were meant to be printed at the receiver on a paper strip, but operators soon realized they could recognize the
letters just by the sound. (Morse code beeping) This sped up the rate at which information could
be sent and received, so Morse code became an
international standard for rapid messaging. (Morse code rapidly beeping) Widely used in the military, maritime communications, and aviation, the universally recognized
distress call, SOS, what does it stand for? Nothing, it just happens
to be easy to send and recognize in Morse code. Morse code revolutionized communication, but in the next century, it would transform a
totally different industry. (mellow jazz music) In the late 1940s, Bernard Silver was an engineering student at Drexel University in Pennsylvania. One day, he overheard the president of a local supermarket chain
asking the engineering dean to find a way to speed
up the checkout process. At that time, cashiers
had to type in each item and its price by hand. A process so tedious and repetitive that many cashiers had developed
carpal tunnel syndrome. Silver told his friend
Norman Joseph Woodland about the problem, and together, they began experimenting. After several false starts,
Woodland moved down to Florida. And one day on the beach, he drew some dots and dashes
of Morse code in the sand, something he was very
familiar with as a Boy Scout. He recalls, "I just extended
the dots and dashes downwards and made narrow lines and
wide lines out of them," and thus the first barcode was created. From this humble beginning, evolved the universal
product code or UPC barcode, capable of storing a simple
string of 12 numbers. It's read by scanning a laser across it and checking how much light is reflected to read the black and white lines, essentially as dots and dashes. A pair of vertical lines are placed at the beginning, middle, and end to ensure the scanner
reads the code properly. The code is divided into
left and right halves. Numbers at either side have their black and white lines flipped so that the scanner can
tell left and right apart, even while reading upside down. When viewed upright, the numbers typically specify
the manufacturer on the left and the product on the right. Manufacturers actually
pay large sums of money to reserve a given amount
of numbers to themselves so that they can exclusively
register their products. (scanner beeps) In this way, the 12 digits
of a barcode uniquely specify every single grocery
item you've ever bought. This jar of Jif peanut butter is identified by the same
12 digits no matter where in the world it's found. And all forms of peanut
butter, smooth, crunchy, stir, no-stir, sugar-free, low-sodium across all brands get
their own unique barcode. Will we ever run out of barcodes? Well, 12 digits could combine
to provide 10 to the 12, that is a trillion different
possible sequences. That should be more than enough, even if companies keep making
stuff like Sour Patch Oreos and Flaming Hot Mountain Dew. But there is a catch. The last digit is not
independent of the others. The barcode creators were aware
that it could get scratched, stained, or tampered with, so
they reserved the last digit to verify that the barcode is complete. Take any barcode and sum the digits at
odd-numbered positions, multiply this result by three, add the digit sum at even-numbered
positions to the result, and take the remainder when
this number is divided by 10. If this remainder is zero,
then the check digit is zero. Otherwise, the check digit
is 10 minus this remainder. (bright electronic music) If a scanner is unable to read
any one digit of the barcode, it can use the final digit to back-calculate what it
must be using this algorithm. But if two digits are damaged,
well, then we're out of luck. In that case, we have to type in the numbers
printed below the barcode. So without that last digit, the number of unique
possibilities is 10 to the 11 or 100 billion options. To date, 1.24 billion
barcodes have been registered, a number that is rising every day, so they won't last us forever. (scanner beeping) But that's not why people began to look for alternatives to barcodes. It was really the amount of information a single
barcode could store. 12 digits by themselves could
only identify one product. But what if someone wanted
to know more information, like where that product came from? (brooding music) In 1986, cattle in the
UK began showing symptoms of a curious brain disease, bovine spongiform encephalopathy, or mad cow disease for short. It is spread when cattle ate feed containing
prions, misfolded proteins, and if people ate beef containing tissue from the brain or spinal
cord of infected cattle, they could contract a
related brain illness that literally turns
your brain into a sponge. Now, as no test could detect
mad cow disease in living cows, millions of cattle were culled. Health officials sought
a method to track sources in imports of beef. But with all the information
this would require for any one piece of beef,
barcodes were insufficient. American inventor David Allais
tried to solve the problem by stacking many barcodes
on top of each other. The result, Code 49,
looked like a bookshelf. This is actually the
predecessor of PDF417, a code often used on
airline boarding passes. But Code 49 didn't solve
the data quantity problem. At best, it multiplied the
information a barcode could carry by a handful. A more efficient way was to extend barcodes
into two dimensions, creating a data matrix. NASA tried this in 1994 with Vericode, used to track and identify
space shuttle parts. This code was read by early digital cameras
rather than lasers and was initially proprietary. Around the same time, Masahiro Hara, an engineer at a Japanese auto
parts manufacturer, DENSO, was getting frustrated at
having to scan multiple barcodes for filling in the same
box of car components. - Hara set out to develop
an alternative on his own. He took inspiration
from an unusual source. - To understand how information
is stored in a QR code, I am going to build one myself in the way Masahiro Hara
first conceptualized it, with a Go board. I have the usual black and white stones. White represents zero.
Black represents one. And we're going to encode the link to our YouTube channel on this board. So the first step is to convert www.youtube.com/veritasium into ones and zeros, and we'll
do that using byte encoding. By encoding uses ASCII, which itself has roots in Morse code. Every character is assigned
a number from 1 to 256. Then we convert the ASCII
decimal into its binary form. Since 256 is 2 to the power of 8, we can use the eight
binary bit combinations to represent all ASCII characters. These eight bits make up
one byte of information. (upbeat techno music) The letter W is assigned
the ASCII decimal 119 or 01110111. Doing the same for all characters in www.youtube.com/veritasium, this is what the string
looks like in binary. This is 26 characters long, so it takes up 26 bytes of information. Now, our board is 25 by 25. This is known as a version 2 QR code, but there are many different sizes, all readable by your phone. Hara's version 1 QR code was 21 by 21, and the largest version
today is 177 by 177. That is large enough to hold
three kilobytes of information. Just 26 of these would've been enough to store all the information
Apollo 11's computer needed to send humans to the moon. (bright music) One programmer (MattKC) even coded
up a computer game, Snake, into a version 40 QR code. (bright music) The region around a QR code must be empty and of uniform color. This is the quiet zone. A distinguishing feature of QR codes is the three square
patterns in the corners. These position squares allow the reader to identify the orientation of the code. Now, almost all QR codes
also have a fourth square in the last corner, but it's smaller and,
hence, trickier to spot. This is the alignment pattern. It's used to rescale the QR code when it's read from varying distances or from crazy oblique angles. The relative size and distance
of the alignment square with respect to the position
squares allows the software to rescale it into a proper square. Next to the position squares
are plain white strips that isolate them from
the rest of the code. And these are timing strips, zebra pedestrian crossings which connect the top left position square with the other two. Every QR code has these
alternating strips. You should look out for them. QR codes of all sizes
visibly look the same, so this tells your phone
which version it is and, therefore, how much data to expect. If there are five alternating
squares, it's version one. If there are nine, it's
version two, and so on. And next to those are format strips that contain rules for
how to scan the code. I'm placing red stones in the
space they occupy for now. There is another feature
every QR code has. This one pixel adjacent to the
bottom right position square, it is always dark. I asked Hara-san if it had
any special significance, but he said no. All of this remaining
space is for data storage. Data inside a QR code always starts at the bottom right corner. Here, the first four
squares carry four bits that specify the data format: 0001 if it's just numbers; 0010 if it's alphanumeric, so
capital letters and numbers; 0100 if it's information stored in bytes; and 1000 for Japanese kanji. The following eight bits are used to indicate
the number of characters in our message. So since we have 26 characters, that should be 00011010. Next, we start arranging our bytes for youtube.com/veritasium starting in eight-bit, two-column cells;. They follow a zigzag pattern that snakes its way to the top left. Within each cell that represents a byte, the most significant bit, corresponding to two to the
seven, is at the bottom right. And the least significant,
or two to the zero, is at the opposite end. 01110111 for w will hence be filled like this. And we'll follow along with the rest. Once we fill in the bytes for www.you, we encounter the alignment
pattern after four bits. To put in the next T, we simply
bypass it and do the same for any of the other
fixed regions of the code. Thus, we keep filling in our data in the same zigzag pattern. After we complete www.youtube.c, the cells start looking less
regular and more Tetris-like. But the way we put the stones byte after byte remains the same. (hopeful music) And there go the last eight
bits for the last letter M. But wait a minute, we've only covered about
half of our QR code. Well, that's because this
whole remaining space is reserved for redundancy. These extra bytes of
error-correction code allow us to reconstruct information
if the QR code is damaged. (menu whooshing) For a fully intact QR code, error correction makes
something else possible, putting a company logo at the center, just like the sponsor
of this video, Saily. Now, I travel a lot. Recently, I was in Germany. Right now, I'm in Australia, and soon, I'm going to the UK. But wherever you are, you
need to have a working phone. Either you pay your home
carrier's hefty roaming fees, or you have to find a place
to buy a local SIM card, put it in the phone, and hope it works. This video's sponsor Saily makes it easy to set up a cell plan and data
in more than 150 countries. You can pick how much data you want and how long you want it for. And it is so much cheaper than roaming. I'm actually going to the UK really soon, and here's how quickly I can
set up an e-SIM with Saily. All I have to do is click on the country, select a plan, and activate the e-SIM. Then when I land, I'll
automatically connect to a local network with no hidden charges. That's it. There's no need
to hunt for public wifi. And you don't have to stand in line at the airport to get a physical SIM. With Saily, you set it up once, and you'll always be connected. And if you find out that your phone isn't compatible with e-SIMs,
you will get a full refund. So to check out Saily for free,
go to Saily.com/veritasium or click the link in the description. Use the code VERITASIUM to get an exclusive 15%
off your first purchase. That's Saily.com/veritasium, or you can scan this handy
QR code to get 15% off. So I wanna thank Saily for sponsoring this part of the video, and now back to building our own QR code. (playful music) QR codes offer four levels
of error correction: low, which can still be read
with 7% of the code missing; medium, which can handle 14%; quartile, 25%; and high, up to 30%. This means a QR code could
still be read properly, even with nearly 1/3 of it missing. Higher levels require more
space for error correction. So knowing how much of the code is error correction is vital. This information is protected in two ways. First, the level of error correction is indicated in the format strip, which is present
identically in two places. The simplest way to avoid errors is to duplicate the information. Here we'll choose the M
level by placing one blue and one yellow stone here at the top left. So what if this part gets damaged? We have a copy in the second format strip starting at the bottom left. The format strip contains three more bits of important information
that we'll get to later. So for now, I'm just gonna
put down three blue stones in both copies. But what about all the
rest of the format strip? Well, this is the second
layer of protection. These other 10 bits are all just designed to correct mistakes in
the first five bits. So how does this work? Let's say I only wanted
to communicate two levels of error correction to you, low or high. If one of the bits flips in transmission to 01 or 10, it's easy to know
that an error has occurred, but no way to know which
the original message was. An easy way to fix this
is to add another bit. So 000 for low, 111 for high. Now, these are at opposite ends of a cube, and hence, they are further apart. If you then receive 011, it's more likely that the
intended message was 111, so it's easy to correct. In this scheme, the
only allowed code words are 000 and 111. The rest act as disallowed buffers to indicate errors in transmission. The allowed code words should
be as far apart as possible. Here, they are three vertices apart. This is known as the Hamming distance, after Richard Hamming who pioneered the field
of error correction. For a Hamming distance of N, you can correct up to N - 1 over 2 errors in a binary string, so one bit flip in the previous example. So back to the five bits
of our format string. If I only wanted to communicate
all zeros or all ones, I could place them at opposite corners of a five-dimensional hypercube. However, our string includes
all 2 to the 5, or 32, combinations of ones and
zeros as valid code words. So to provide buffers like before, we can extend the 5-bit
string into a 15-bit string. And now the 32 valid code words are each separated by seven vertices, or a Hamming distance of seven, which means we can correct
up to three bit flip errors. The easiest way to do this
is using a lookup table. The table takes a slightly misread vertex and finds the closest valid vertex, likely the intended code word. But for our main QR code data, we need a far more efficient scheme, one that doesn't require lookup tables or doubling or tripling our data size. (bright music) Let's say I wanna send you a message that is the four numbers 1, -2, 3 and 5. If I just send these numbers, one of them could get
corrupted in transmission, and you wouldn't know
that an error had occurred or which digit was wrong. So before I send the message,
we come up with a plan. First, instead of sending you
four numbers, I will send six. The first four are my actual message, and the last two, A and B, will help you check if
there were any errors. Now, I want you to treat these six numbers as the coefficients of a
degree-five polynomial, and I will pick the values of A and B so that this polynomial
could also be written in the form of a degree-three
polynomial, call it q of x, times (x - 1)(x - 2). Now, we could set these last two terms to be x minus any number, but for simplicity, let's
say we pick one and two. That way, when you receive my polynomial, you know that if you
plug in x = 1 or x = 2, you should get zero for both because that's how I
constructed the polynomial. And if you don't get zero, you know there has been
an error in transmission. They are called syndromes. Which is an apt term since syndromes are defined as a group of signs that occur together and characterize a particular abnormality. If the message polynomial is not 0 at any of the syndrome values then there is an error in the code. So how do I find the values
of 'a' and 'b' in our example? Well, I take the
polynomial without 'a' and 'b' and divide it by x - 1, x - 2. I get a degree-three polynomial, which is what I want, but
there is also a remainder of 37x - 30. So I can move this to the left-hand side. For the polynomial
to take the form I want, 'a' must be negative 37 and 'b' positive 30. So I send the message 1, -2, 3, 5, - 37, and 30. You can plug in x = 1 and x = 2, and if you get zero for both, you know the message was sent correctly. But what if there was an
error in transmission? Say at position four, the
number has changed to a six. Well, now if you evaluate at x = 1 and 2, the polynomial is no longer zero. To figure out where the error occurred, one at a time, you set each
coefficient to be a variable. Then find the value of that variable, setting the polynomial
equal to 0 at x = 1. You repeat this for x = 2. And what you find is that
the two values are different. This indicates that the second coefficient was not the error. You find the same for all
the other coefficients, except when you reach the
one where the error occurred. Here, not only are the two values equal, they're also equal to the
originally transmitted number. That is, five. So this method allows us both to check and correct errors with only a
modest increase in data size. (inquisitive music) This is a toy example of a Reed-Solomon error correcting code, developed by mathematicians Irving S. Reed and Gustav Solomon in 1960. (upbeat music) (QR code whooshing) The job of decoding Reed-Solomon
codes in a brute-force way, as we described, can
quickly get intensive. In fact, as the Voyager spacecraft floated into the outer solar system, NASA engineers knew their
signal-to-noise ratio would get incredibly small. But the promise of
Reed-Solomon codes was such that they put an experimental
encoder in before launch, wagering that that smarter
encoding algorithms would follow in the next decade. And that's exactly what happened. To this day, we can make out Voyager's
ever-faintening whispers, thanks to Reed-Solomon codes. These codes also ensure that your old CDs or DVDs can still play your favorite songs and movies despite multiple scratches. And they are the reason why QR codes still work when damaged. In a QR code, the entire data, starting from the data
type, character length byte, our message bytes, and final padding, are laid out in a line and converted back into ASCII decimals. Fitting a high-degree
polynomial using these can easily make the coefficients blow up. Hence, Reed-Solomon encoding
uses finite-field arithmetic, Galois fields, to obtain
the error correcting terms. These, converted back into binary, are used to fill out
the rest of the QR code. (bright music) And there we have our complete QR code. But why can't we scan it yet? See how these regions here
appear uniformly white and black? Well, sometimes the encoded
data can insert plain patterns and blank spaces just by chance. These can confound the readers which expect to see a noisy checkerboard. They think maybe it's a big damage patch, or maybe it's not a QR code at all. But there is a way to fix this. Remember the three blue stones I put in for masking in the format string? Well, they specify one of eight ways to reshuffle the appearance
of our QR code pixels to make them seem truly jumbled. Now, this particular mask says, flip the pixels so white becomes black and black becomes white for
every third column of data. But this does not apply to the functional elements of the code. They remain unchanged. (bars whooshing) (QR code beeping) The QR code standard specifies the use of eight masking patterns. In principle, when combined
with the correct masking bits, all eight forms of the code are readable, which is why some QR code generators will return different looking codes for the same input string. But which one works best? Well, every continuous
or bad patch adds points, and each mask gets assigned a score. The mask with the lowest
score at the end wins. It's easiest for any reader to scan. (confetti cannon pops) (audience applauds) For our handmade QR code, I'm going to use the simplest mask. And now we have a working QR code. Try it out. (bright music) And moment of truth. - Ahh, it worked. You know, going through this
exercise made me realize again why I hate QR codes. They're not meant for people. I made all kinds of mistakes while I was trying to put down
these black and white stones. This was really hard to get this perfect. I guess it didn't have to be perfect, but it had to be close enough, all right. Initially, QR codes had
only industrial uses. But it wasn't long before the value of their data storage
capacity was realized. In 2002, mad cow disease
resurfaced in the UK. 179 people died from
eating contaminated beef, and people panicked. They wanted to know exactly
where their meat was coming from and how it was stored before
it reached the supermarket. This time, the QR code
was available to help. It was one of the first instances where the curious checkerboards started to appear in common use. But why are QR codes so successful? There are plenty of other
2D matrix codes out there. Well, one reason is that
DENSO Wave decided not to exercise patent rights on QR codes. We made the patent open to everyone, which made the QR code so popular. DENSO instead opted to monetize
and sell QR code scanners. Of course, with the rise of smartphones, most people would soon
carry a QR code scanner in their pocket. But initially, QR code reading apps were third-party and rather niche. But then in 2017, Android and Apple built
QR code readers right into their camera apps, so the use of these codes took off. The COVID-19 pandemic also gave
QR codes a boost worldwide. Suddenly, restaurants and
vendors wanted a contactless way to hand out menus and product information. Contactless payment
using QR codes took off in India and China. Today, India sees over 12 billion QR code-enabled
transactions per month. QR codes also proved handy
in storing vaccine records and personal health
information in phone wallets. But their enormous spread
has also created problems. A question about QR code safety: In recent years, some
scammers have used QR codes to try to defraud people who read them. Do you have concerns about these uses? So as with anything on the internet, you have to pay extra attention to safety. Check where a scanned
QR code is taking you before actually clicking on a link. So I wanna know about
the future of QR codes. What is next for QR codes? Now, if the possibility of running out of UPC barcodes is remote, for QR codes, it is impossible. The number of unique version 1 QR codes using the lowest redundancy
level is 2 to the 152. This is about 10 times the total number of legal chessboard configurations, which is also why a random distribution of pixels filled into a QR code pattern generally cannot be
interpreted as a message. You have scanned countless QR codes, and when you scan your next one, you'll have a better idea of how it works. But have you ever thought about
what QR itself stands for? What is your favorite
application of the QR code? (space tones beeping)"""
summaries = []

for chunk in chunk_text(text):
    chunk_length = len(chunk.split())

    if 1 <= chunk_length <= 150:
        min_len = 1
        max_len = 60
    elif 151 <= chunk_length <= 300:
        min_len = 60
        max_len = 120
    else:
        min_len = 100
        max_len = 250

    summary = summarizer(chunk, max_length=max_len, min_length=min_len, do_sample=False)
    summaries.append(summary[0]['summary_text'])

final_summary = ' '.join(summaries)

print(final_summary)

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def save_to_word(summary, filename='summary.docx'):
    doc = Document()
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    doc.save(filename)
    print(f'Summary saved to Word file: {filename}')

# Save summary to a PDF file
def save_to_pdf(summary, filename='summary.pdf'):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 100, 'Summary')
    c.drawString(100, height - 120, summary)
    c.save()
    print(f'Summary saved to PDF file: {filename}')

# Save the summary
save_to_word(final_summary)
save_to_pdf(final_summary)
